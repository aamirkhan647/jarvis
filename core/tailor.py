# core/tailor.py
"""
Tailor module: create a simple targeted DOCX from the original resume + job description.
Optional OpenAI integration commented for safety.
"""

from docx import Document
from docx.shared import Pt
from datetime import datetime
import re
import logging

logger = logging.getLogger(__name__)


def synthesize_summary(resume_text, job_text):
    # Simple keyword-based short summary fallback
    jd_words = re.findall(r"\b[a-zA-Z0-9_+-/.]+\b", (job_text or "").lower())
    top = []
    freq = {}
    for w in jd_words:
        if len(w) > 2:
            freq[w] = freq.get(w, 0) + 1
    sorted_k = sorted(freq.items(), key=lambda x: x[1], reverse=True)
    top = [k for k, _ in sorted_k[:6]]
    if not top:
        return "Experienced professional seeking a role that leverages technical skills and domain knowledge."
    return (
        "Experienced with "
        + ", ".join(top[:5])
        + " — focused on delivering production-grade software."
    )


def extract_skills_from_jd(job_text, top_n=10):
    words = re.findall(r"\b[a-zA-Z0-9_+.-]+\b", (job_text or "").lower())
    freq = {}
    for w in words:
        if len(w) > 1:
            freq[w] = freq.get(w, 0) + 1
    keys = sorted(freq.items(), key=lambda x: x[1], reverse=True)
    return [k for k, _ in keys[:top_n]]


def create_docx(resume_text, job_text, output_path, name="Your Name"):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Header
    h = doc.add_paragraph()
    run = h.add_run(name)
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph("Generated: " + datetime.now().strftime("%Y-%m-%d %H:%M"))

    # Targeted Summary
    doc.add_heading("Profile", level=2)
    doc.add_paragraph(synthesize_summary(resume_text, job_text))

    # Skills
    doc.add_heading("Skills", level=2)
    skills = extract_skills_from_jd(job_text, top_n=12)
    if skills:
        doc.add_paragraph(", ".join(skills))
    else:
        doc.add_paragraph("See experience for relevant skills.")

    # Experience / Education (dump resume text)
    doc.add_heading("Experience & Education", level=2)
    for p in re.split(r"\n\s*\n", (resume_text or "")):
        if p.strip():
            doc.add_paragraph(p.strip()[:1500])

    doc.save(output_path)
    logger.info("Saved tailored docx to %s", output_path)
    return output_path
